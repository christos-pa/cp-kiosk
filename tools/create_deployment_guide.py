from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"C:\Users\paraskevopoulos.chri\Documents\New project\deployment\CP Kiosk Deployment Guide v0.2.8.docx"

BLUE = "147D72"
DARK = "152131"
MUTED = "5F6B76"
LIGHT_BLUE = "EAF5F3"
LIGHT_GRAY = "F3F5F7"
CAUTION = "FFF4D6"
RED = "A12622"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=140, start=180, bottom=140, end=180):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph_box(paragraph, fill, border_color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)

    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for side in ("top", "left", "bottom", "right"):
        edge = p_bdr.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            p_bdr.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "8" if side == "left" else "2")
        edge.set(qn("w:space"), "8")
        edge.set(qn("w:color"), border_color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_fixed_table_width(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run(run, size=None, bold=None, color=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Guide Note" not in styles:
        note = styles.add_style("Guide Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = styles["Guide Note"]
    note.font.name = "Calibri"
    note.font.size = Pt(10.5)
    note.font.color.rgb = RGBColor.from_string(DARK)
    note.paragraph_format.space_before = Pt(0)
    note.paragraph_format.space_after = Pt(4)
    note.paragraph_format.line_spacing = 1.15


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("CP Kiosk | Deployment Guide")
    set_run(run, size=9, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    add_page_number(p)


def add_callout(doc, title, lines, fill=LIGHT_BLUE, title_color=BLUE):
    p = doc.add_paragraph()
    p.style = "Guide Note"
    p.paragraph_format.left_indent = Inches(0.06)
    p.paragraph_format.right_indent = Inches(0.06)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    set_paragraph_box(p, fill, title_color)
    r = p.add_run(title)
    set_run(r, size=11, bold=True, color=title_color)
    for line in lines:
        p.add_run().add_break()
        if isinstance(line, tuple):
            label, value = line
            r = p.add_run(label)
            set_run(r, size=10.5, bold=True, color=DARK)
            r = p.add_run(value)
            set_run(r, size=10.5, color=DARK)
        else:
            r = p.add_run(line)
            set_run(r, size=10.5, color=DARK)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_run(r)
    else:
        r = p.add_run(text)
        set_run(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run(r)
    return p


def add_option_table(doc):
    rows = [
        ("Kiosk active", "Turns kiosk enforcement on or off. Fresh installations start disabled."),
        ("Reload current page", "Refreshes the page currently displayed."),
        ("Load start page", "Loads the URL saved in Kiosk settings."),
        ("Clear browser data", "Deletes cookies, cache and browser storage. This normally signs the website out."),
        ("Back", "Moves to the previous browser page, when available."),
        ("Forward", "Moves to the next browser page, when available."),
        ("Settings", "Changes the website URL and four-digit administration PIN."),
        ("Exit kiosk", "Temporarily returns to Android Home. It does not switch off Kiosk active."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_fixed_table_width(table, [2700, 6660])
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(("Option", "Purpose")):
        cell = header.cells[idx]
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_run(r, bold=True, color=DARK)
    for option, purpose in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((option, purpose)):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run(r, bold=(idx == 0))
    set_fixed_table_width(table, [2700, 6660])


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    configure_styles(doc)
    add_header_footer(section)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("DEPLOYMENT")
    set_run(r, size=11, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("CP Kiosk Deployment Guide")
    set_run(r, size=26, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Version 0.2.8 | Strong kiosk prioritised | Soft kiosk fallback")
    set_run(r, size=12, color=MUTED)

    add_callout(
        doc,
        "Easiest APK access: Google Drive",
        [
            "Sign in to the following Google account on an authorised office device and open the shared CP Kiosk deployment folder.",
            ("Account: ", "[deployment account]"),
            ("Password: ", "[obtain from the authorised administrator]"),
            "Use only the APK that matches the intended deployment type and version.",
            "Do not place deployment credentials directly in this document or the source repository.",
        ],
        fill=CAUTION,
        title_color="7A5A00",
    )

    doc.add_heading("Deployment files", level=1)
    add_bullet(doc, "Strong: cp-kiosk-strong-v0.2.8-google-disabled-default.apk", "Strong:")
    add_bullet(doc, "Soft: cp-kiosk-soft-v0.2.8-google-disabled-default.apk", "Soft:")
    add_bullet(doc, "Fresh-install website: https://www.google.com", "Fresh-install website:")
    add_bullet(doc, "Fresh-install kiosk state: Disabled", "Fresh-install kiosk state:")
    add_bullet(doc, "Default administration PIN: 4711", "Default administration PIN:")

    add_callout(
        doc,
        "Important",
        [
            "Installing an update over an existing CP Kiosk installation normally preserves its saved URL, kiosk state, cookies and website session. The Google URL and disabled state are defaults for a fresh installation or cleared app data."
        ],
        fill=LIGHT_GRAY,
    )

    doc.add_heading("Choose the deployment type", level=1)
    p = doc.add_paragraph()
    r = p.add_run("Use Strong kiosk whenever the tablet can be prepared correctly. ")
    set_run(r, bold=True)
    r = p.add_run(
        "Use Soft kiosk when the tablet is already deployed, cannot be reset, or cannot meet the device-owner conditions."
    )
    set_run(r)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_fixed_table_width(table, [1900, 3730, 3730])
    set_repeat_table_header(table.rows[0])
    for idx, text in enumerate(("Version", "Best for", "Main limitation")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "E8EEF5")
        r = cell.paragraphs[0].add_run(text)
        set_run(r, bold=True)
    for values in (
        ("Strong", "New or reset tablets that can be assigned to CP Kiosk as Android device owner.", "Requires a clean device-owner setup over USB/ADB."),
        ("Soft", "Existing tablets that must remain configured and cannot be reset.", "Uses normal app/screen-pinning behaviour and is easier to escape than Strong kiosk."),
    ):
        cells = table.add_row().cells
        for idx, text in enumerate(values):
            r = cells[idx].paragraphs[0].add_run(text)
            set_run(r, bold=(idx == 0))
    set_fixed_table_width(table, [1900, 3730, 3730])

    doc.add_heading("Strong kiosk - recommended", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "Strong kiosk uses Android's official device-owner and lock-task controls. "
        "It does not require Samsung Knox, but it must be provisioned before normal accounts and usage are established."
    )
    set_run(r)

    doc.add_heading("Strong kiosk conditions", level=2)
    add_bullet(doc, "The tablet should be brand new or factory reset.")
    add_bullet(doc, "No Google, Samsung, Microsoft, Exchange or other user accounts may remain on the tablet.")
    add_bullet(doc, "Initial setup must be completed without restoring a backup.")
    add_bullet(doc, "Developer Options and USB debugging must be enabled.")
    add_bullet(doc, "The tablet must be connected to the deployment PC and the USB debugging prompt accepted.")
    add_bullet(doc, "CP Kiosk must be assigned as Android device owner using ADB.")

    add_callout(
        doc,
        "If Android refuses device-owner setup",
        [
            "The usual reason is that one or more accounts already exist on the tablet. Removing accounts can sometimes be sufficient, but a factory reset is the safest and most reliable route. If the tablet cannot be reset, deploy the Soft kiosk version instead."
        ],
        fill=CAUTION,
        title_color="7A5A00",
    )

    doc.add_heading("Strong kiosk installation", level=2)
    for step in (
        "Factory reset the tablet.",
        "Complete only the minimum Android setup. Skip Google account, Samsung account, backup restore and optional personalisation.",
        "Open Settings > About tablet > Software information and tap Build number seven times to enable Developer Options.",
        "Open Settings > Developer options and enable USB debugging.",
        "Connect the tablet to the deployment PC by USB and select Allow on the USB debugging prompt.",
        "Install the Strong v0.2.8 APK and assign com.example.platekiosk/.KioskDeviceAdminReceiver as device owner using ADB.",
        "Launch CP Kiosk. It starts with Kiosk active switched off.",
        "Swipe from the left edge to the right, then turn Kiosk active on. Once active, CP Kiosk enters permanent strong lock-task mode.",
    ):
        add_number(doc, step)

    add_callout(
        doc,
        "Strong kiosk result",
        [
            "The app becomes the managed device owner, Home/Recent-app escape routes are restricted, CP Kiosk can return automatically, and Android may block normal uninstall until device-owner status is removed or the tablet is reset."
        ],
        fill=LIGHT_BLUE,
    )

    doc.add_heading("Soft kiosk - fallback", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "Use Soft kiosk when the tablet is already configured, contains accounts, is deployed on site, or cannot be factory reset."
    )
    set_run(r)

    doc.add_heading("Soft kiosk installation", level=2)
    for step in (
        "Open Google Drive using the deployment account shown at the top of this guide.",
        "Download cp-kiosk-soft-v0.2.8-google-disabled-default.apk.",
        "Open the APK. If Android asks, allow the selected Drive/browser/file app to install unknown applications.",
        "Install or update CP Kiosk. An in-place update using the same signed APK should preserve app data and cookies.",
        "Set CP Kiosk as the Home app when prompted and keep Android screen pinning enabled where required.",
        "Open CP Kiosk. Swipe from the left edge to the right and turn Kiosk active on when the tablet is ready to enter service.",
    ):
        add_number(doc, step)

    add_callout(
        doc,
        "Soft kiosk limitation",
        [
            "Soft kiosk is practical for existing tablets, but it does not have the same Android device-owner authority as Strong kiosk. Some system bars, screen-pinning messages or escape routes may still be controlled by Android."
        ],
        fill=LIGHT_GRAY,
    )

    doc.add_heading("Kiosk administration", level=1)
    doc.add_heading("Open the administration menu", level=2)
    add_bullet(doc, "Swipe from the left edge of the kiosk screen toward the right.")
    add_bullet(doc, "If Kiosk active is ON, enter the four-digit administration PIN. Default: 4711.")
    add_bullet(doc, "If Kiosk active is OFF, the administration menu opens directly without requesting a PIN.")

    add_callout(
        doc,
        "Launching the kiosk",
        [
            "A fresh installation starts disabled. Before putting the tablet into service, open Kiosk administration and switch Kiosk active ON."
        ],
        fill=LIGHT_BLUE,
    )

    doc.add_heading("Administration options", level=2)
    add_option_table(doc)

    doc.add_heading("Change the URL or PIN", level=1)
    for step in (
        "Open Kiosk administration using the left-to-right swipe.",
        "Select Settings.",
        "Under Website, enter the required URL.",
        "Under Security, review or replace the four-digit PIN. The default PIN is 4711.",
        "Select Save. CP Kiosk loads the saved start page.",
    ):
        add_number(doc, step)

    add_callout(
        doc,
        "Cookie and session guidance",
        [
            "CP Kiosk is designed to preserve browser cookies and website sessions across normal app exits, restarts and same-signature APK updates.",
            "Do not select Clear browser data unless you intentionally want to remove cookies and sign the website out.",
            "Uninstalling the app, clearing Android app data, changing the signing certificate or factory resetting the tablet can remove the stored website session.",
        ],
        fill=CAUTION,
        title_color="7A5A00",
    )

    doc.add_heading("Quick deployment checklist", level=1)
    for item in (
        "Correct APK type selected: Strong preferred, Soft fallback.",
        "Installed version confirmed as 0.2.8.",
        "Website URL checked in Settings.",
        "Administration PIN confirmed; default is 4711.",
        "Kiosk active switched ON before the tablet enters service.",
        "Website login/cookie session tested.",
        "For Strong kiosk, device-owner status and locked mode verified.",
    ):
        add_bullet(doc, item)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Internal deployment guide - CP Kiosk v0.2.8")
    set_run(r, size=9, color=MUTED, italic=True)

    doc.core_properties.title = "CP Kiosk Deployment Guide v0.2.8"
    doc.core_properties.subject = "Strong and soft Android kiosk deployment guidance"
    doc.core_properties.author = "CP Kiosk Operations"
    doc.core_properties.keywords = "CP Kiosk, Android, Strong Kiosk, Soft Kiosk, Deployment"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
